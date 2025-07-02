#!/usr/bin/env python3
"""
Advanced Knowledge Extraction System for K-Array
Estrae e processa tutti i tipi di contenuti da k-array.com con focus su keywords.
"""

import os
import sys
import json
import logging
import asyncio
import aiohttp
import aiofiles
from pathlib import Path
from typing import Dict, List, Set, Optional, Tuple, Any
from dataclasses import dataclass, asdict
from datetime import datetime
import hashlib
import re
from urllib.parse import urljoin, urlparse
import mimetypes

# Import libraries per processing diversi formati
try:
    import PyPDF2
    from pdfplumber import PDF
except ImportError:
    PyPDF2 = None
    PDF = None

try:
    import openpyxl
    from openpyxl import load_workbook
except ImportError:
    openpyxl = None

try:
    import docx
    from docx import Document
except ImportError:
    docx = None

try:
    import requests
    from bs4 import BeautifulSoup
except ImportError:
    requests = None
    BeautifulSoup = None

# Setup logging
logging.basicConfig(level=logging.INFO)
logger = logging.getLogger(__name__)

@dataclass
class ExtractedDocument:
    """Struttura per documenti estratti con metadata avanzati."""
    id: str
    content: str
    title: str
    source_url: str
    file_type: str
    category: str
    keywords: List[str]
    product_tags: List[str]
    technical_specs: Dict[str, Any]
    confidence_score: float
    extraction_timestamp: str
    file_hash: str
    size_bytes: int
    language: str = "it"
    
    def to_jsonl(self) -> str:
        """Converte in formato JSONL per storage."""
        data = asdict(self)
        return json.dumps(data, ensure_ascii=False)

class KeywordExtractor:
    """Estrattore di keywords avanzato specifico per K-Array."""
    
    def __init__(self):
        # Keywords specifiche K-Array per categorizzazione
        self.product_keywords = {
            'k-framework': ['k-framework', 'kframework', 'framework', 'software', 'acoustical', 'simulation', '3d'],
            'thunder': ['thunder', 'subwoofer', 'bass', 'ks5', 'ks3', 'low frequency'],
            'mugello': ['mugello', 'line array', 'kh2', 'kh3', 'kh5', 'live sound'],
            'vyper': ['vyper', 'aluminum', 'flat', 'line array', 'elegant'],
            'anakonda': ['anakonda', 'kan200', 'flexible', 'problem solver'],
            'kobra': ['kobra', 'stainless steel', 'passive', 'pure array'],
            'python': ['python', 'stainless steel', '3.15 inch', 'drivers'],
            'dragon': ['dragon', 'point source', 'high pressure', 'powerful'],
            'turtle': ['turtle', 'monitor', 'wedge', 'low-profile'],
            'capture': ['capture', 'microphone', 'kmc20', 'kmc50', 'cardioid'],
            'duetto': ['duetto', 'earbuds', 'kd6t', 'kd6b', 'lifestyle'],
            'azimut': ['azimut', 'complete solution', 'minimalist', 'discreet'],
            'pinnacle': ['pinnacle', 'powered', 'lightweight', 'portable'],
            'prisma': ['prisma', 'matrix', 'processing', 'platform'],
            'kommander': ['kommander', 'amplifier', 'ka104', 'ka208', 'dante'],
        }
        
        self.technical_keywords = {
            'audio': ['spl', 'frequency', 'impedance', 'watt', 'ohm', 'db', 'hz', 'khz'],
            'connectivity': ['dante', 'ethernet', 'wifi', 'bluetooth', 'usb', 'xlr', 'rca'],
            'features': ['beam steering', 'ebs', 'pure array', 'cardioid', 'dsp', 'preset'],
            'applications': ['live sound', 'installed sound', 'touring', 'stadium', 'yacht', 'venue'],
            'materials': ['aluminum', 'stainless steel', 'weather resistant', 'marine grade']
        }
        
        self.category_keywords = {
            'software': ['software', 'k-framework', 'k-monitor', 'web app', 'k-connect', 'api'],
            'speakers': ['speaker', 'line array', 'point source', 'subwoofer', 'monitor'],
            'amplifiers': ['amplifier', 'kommander', 'powered', 'processing', 'matrix'],
            'microphones': ['microphone', 'capture', 'cardioid', 'array'],
            'accessories': ['accessory', 'mounting', 'bracket', 'cable', 'adapter'],
            'lifestyle': ['lifestyle', 'earbuds', 'duetto', 'personal', 'mobile']
        }
    
    def extract_keywords(self, text: str, title: str = "") -> Tuple[List[str], List[str], str]:
        """
        Estrae keywords, product tags e categoria da un testo.
        Returns: (keywords, product_tags, category)
        """
        text_lower = text.lower()
        title_lower = title.lower()
        combined_text = f"{title_lower} {text_lower}"
        
        # Estrai keywords generali
        keywords = set()
        product_tags = set()
        
        # Product detection
        for product, terms in self.product_keywords.items():
            for term in terms:
                if term in combined_text:
                    product_tags.add(product)
                    keywords.add(term)
        
        # Technical keywords
        for category, terms in self.technical_keywords.items():
            for term in terms:
                if term in combined_text:
                    keywords.add(term)
        
        # Determina categoria principale
        category_scores = {}
        for cat, terms in self.category_keywords.items():
            score = sum(1 for term in terms if term in combined_text)
            if score > 0:
                category_scores[cat] = score
        
        primary_category = max(category_scores.items(), key=lambda x: x[1])[0] if category_scores else "general"
        
        # Aggiungi keywords specifiche da pattern
        keywords.update(self._extract_model_numbers(combined_text))
        keywords.update(self._extract_technical_specs(combined_text))
        
        return list(keywords), list(product_tags), primary_category
    
    def _extract_model_numbers(self, text: str) -> Set[str]:
        """Estrae numeri di modello e codici prodotto."""
        patterns = [
            r'\bk[a-z]*[\-]?[0-9]+[a-z]*\b',  # K-codes: KA104, K-Framework, etc.
            r'\b[a-z]+-[a-z]+[0-9]*\b',       # Product codes: Thunder-KS5
            r'\b[a-z]{2,}[0-9]+[a-z]*\b',     # Model codes: KMC20, KAN200
        ]
        
        models = set()
        for pattern in patterns:
            matches = re.findall(pattern, text, re.IGNORECASE)
            models.update(matches)
        
        return models
    
    def _extract_technical_specs(self, text: str) -> Set[str]:
        """Estrae specifiche tecniche come frequenze, potenze, etc."""
        patterns = [
            r'\b\d+\s*hz\b',           # Frequency: 20 Hz
            r'\b\d+\s*khz\b',          # Frequency: 20 kHz  
            r'\b\d+\s*w\b',            # Power: 100 W
            r'\b\d+\s*watt\b',         # Power: 100 watt
            r'\b\d+\s*ohm\b',          # Impedance: 8 ohm
            r'\b\d+\s*db\b',           # SPL: 120 dB
            r'\b\d+\"\b',              # Size: 4"
            r'\b\d+\s*inch\b',         # Size: 4 inch
        ]
        
        specs = set()
        for pattern in patterns:
            matches = re.findall(pattern, text, re.IGNORECASE)
            specs.update(matches)
        
        return specs

class AdvancedExtractor:
    """Sistema di estrazione avanzata multi-formato."""
    
    def __init__(self, base_url: str = "https://www.k-array.com"):
        self.base_url = base_url
        self.keyword_extractor = KeywordExtractor()
        self.session = None
        self.processed_urls = set()
        self.extracted_documents = []
        self.download_dir = Path("downloads")
        self.download_dir.mkdir(exist_ok=True)
        
        # Statistics
        self.stats = {
            'total_pages': 0,
            'pdfs_processed': 0,
            'excel_processed': 0,
            'word_processed': 0,
            'videos_found': 0,
            'errors': 0,
            'duplicates_skipped': 0
        }
    
    async def __aenter__(self):
        """Async context manager entry."""
        connector = aiohttp.TCPConnector(limit=10, limit_per_host=5)
        timeout = aiohttp.ClientTimeout(total=30, connect=10)
        self.session = aiohttp.ClientSession(
            connector=connector,
            timeout=timeout,
            headers={
                'User-Agent': 'K-Array Knowledge Extractor 1.0',
                'Accept': 'text/html,application/xhtml+xml,application/xml;q=0.9,*/*;q=0.8'
            }
        )
        return self
    
    async def __aexit__(self, exc_type, exc_val, exc_tb):
        """Async context manager exit."""
        if self.session:
            await self.session.close()
    
    async def extract_all(self, start_urls: List[str] = None) -> List[ExtractedDocument]:
        """
        Estrae tutti i contenuti da K-Array con approccio sistematico.
        """
        if start_urls is None:
            start_urls = [
                f"{self.base_url}/en",
                f"{self.base_url}/en/products",
                f"{self.base_url}/en/software", 
                f"{self.base_url}/en/support",
                f"{self.base_url}/en/downloads",
                f"{self.base_url}/en/applications"
            ]
        
        logger.info(f"Starting advanced extraction from {len(start_urls)} URLs")
        
        # Phase 1: Discover all pages and downloadable content
        all_urls = await self._discover_all_urls(start_urls)
        logger.info(f"Discovered {len(all_urls)} unique URLs")
        
        # Phase 2: Extract content from each URL
        for url in all_urls:
            try:
                await self._process_url(url)
                await asyncio.sleep(0.5)  # Rate limiting
            except Exception as e:
                logger.error(f"Error processing {url}: {e}")
                self.stats['errors'] += 1
        
        # Phase 3: Process downloaded files
        await self._process_downloaded_files()
        
        # Phase 4: Quality validation and deduplication
        self._validate_and_deduplicate()
        
        logger.info(f"Extraction complete. Stats: {self.stats}")
        logger.info(f"Total documents extracted: {len(self.extracted_documents)}")
        
        return self.extracted_documents
    
    async def _discover_all_urls(self, start_urls: List[str]) -> Set[str]:
        """Scopre tutti gli URL del sito tramite crawling sistematico."""
        discovered = set(start_urls)
        to_crawl = list(start_urls)
        crawled = set()
        
        while to_crawl:
            current_batch = to_crawl[:10]  # Process in batches
            to_crawl = to_crawl[10:]
            
            tasks = []
            for url in current_batch:
                if url not in crawled:
                    tasks.append(self._extract_links_from_page(url))
                    crawled.add(url)
            
            if tasks:
                results = await asyncio.gather(*tasks, return_exceptions=True)
                
                for result in results:
                    if isinstance(result, list):
                        for new_url in result:
                            if new_url not in discovered and self._is_relevant_url(new_url):
                                discovered.add(new_url)
                                if new_url not in crawled:
                                    to_crawl.append(new_url)
        
        return discovered
    
    def _is_relevant_url(self, url: str) -> bool:
        """Determina se un URL è rilevante per l'estrazione."""
        parsed = urlparse(url)
        
        # Must be from k-array.com
        if 'k-array.com' not in parsed.netloc:
            return False
        
        # Skip irrelevant paths
        skip_patterns = [
            '/cart', '/checkout', '/login', '/register', '/account',
            '/search', '/admin', '/api/', '/_',
            '.css', '.js', '.ico', '.png', '.jpg', '.gif', '.svg'
        ]
        
        for pattern in skip_patterns:
            if pattern in url.lower():
                return False
        
        return True
    
    async def _extract_links_from_page(self, url: str) -> List[str]:
        """Estrae tutti i link da una pagina."""
        try:
            async with self.session.get(url) as response:
                if response.status != 200:
                    return []
                
                content = await response.text()
                soup = BeautifulSoup(content, 'html.parser')
                
                links = []
                for tag in soup.find_all(['a', 'link']):
                    href = tag.get('href')
                    if href:
                        absolute_url = urljoin(url, href)
                        links.append(absolute_url)
                
                # Find downloadable files
                for tag in soup.find_all(['a']):
                    href = tag.get('href', '')
                    if any(ext in href.lower() for ext in ['.pdf', '.docx', '.xlsx', '.xls', '.doc', '.zip']):
                        absolute_url = urljoin(url, href)
                        links.append(absolute_url)
                
                return links
                
        except Exception as e:
            logger.error(f"Error extracting links from {url}: {e}")
            return []
    
    async def _process_url(self, url: str):
        """Processa un singolo URL."""
        if url in self.processed_urls:
            self.stats['duplicates_skipped'] += 1
            return
        
        self.processed_urls.add(url)
        
        try:
            async with self.session.get(url) as response:
                if response.status != 200:
                    return
                
                content_type = response.headers.get('content-type', '').lower()
                
                if 'text/html' in content_type:
                    await self._process_html_page(url, await response.text())
                    self.stats['total_pages'] += 1
                    
                elif 'application/pdf' in content_type:
                    await self._download_and_queue_file(url, response)
                    
                elif any(mime in content_type for mime in ['application/vnd.openxmlformats', 'application/vnd.ms-excel']):
                    await self._download_and_queue_file(url, response)
                    
                elif 'application/vnd.openxmlformats-officedocument' in content_type:
                    await self._download_and_queue_file(url, response)
        
        except Exception as e:
            logger.error(f"Error processing URL {url}: {e}")
            self.stats['errors'] += 1
    
    async def _process_html_page(self, url: str, html_content: str):
        """Processa una pagina HTML estraendo contenuto e metadata."""
        soup = BeautifulSoup(html_content, 'html.parser')
        
        # Remove navigation, footer, sidebar content
        for element in soup.find_all(['nav', 'footer', 'aside', 'header']):
            element.decompose()
        
        # Extract title
        title_tag = soup.find('title')
        title = title_tag.text.strip() if title_tag else ""
        
        # Extract main content
        main_content = ""
        
        # Try to find main content areas
        content_selectors = [
            'main', '.main-content', '#main', '.content',
            '.product-description', '.product-info', 
            'article', '.article-content'
        ]
        
        for selector in content_selectors:
            elements = soup.select(selector)
            if elements:
                main_content = " ".join(elem.get_text(strip=True, separator=' ') for elem in elements)
                break
        
        # Fallback: extract all visible text
        if not main_content:
            main_content = soup.get_text(strip=True, separator=' ')
        
        # Clean content
        main_content = self._clean_text(main_content)
        
        if len(main_content) < 100:  # Skip pages with little content
            return
        
        # Extract keywords and categorize
        keywords, product_tags, category = self.keyword_extractor.extract_keywords(main_content, title)
        
        # Calculate confidence score
        confidence = self._calculate_confidence_score(main_content, title, keywords)
        
        # Create document
        doc = ExtractedDocument(
            id=self._generate_id(url),
            content=main_content,
            title=title,
            source_url=url,
            file_type="html",
            category=category,
            keywords=keywords,
            product_tags=product_tags,
            technical_specs=self._extract_technical_specs(main_content),
            confidence_score=confidence,
            extraction_timestamp=datetime.now().isoformat(),
            file_hash=hashlib.md5(main_content.encode()).hexdigest(),
            size_bytes=len(main_content.encode()),
            language="en" if "/en/" in url else "it"
        )
        
        self.extracted_documents.append(doc)
        logger.info(f"Extracted: {title[:50]}... (Category: {category}, Keywords: {len(keywords)})")
    
    async def _download_and_queue_file(self, url: str, response):
        """Scarica un file per processamento offline."""
        try:
            filename = Path(urlparse(url).path).name
            if not filename:
                filename = f"download_{hashlib.md5(url.encode()).hexdigest()[:8]}"
            
            file_path = self.download_dir / filename
            
            # Download file
            content = await response.read()
            async with aiofiles.open(file_path, 'wb') as f:
                await f.write(content)
            
            logger.info(f"Downloaded: {filename} ({len(content)} bytes)")
            
        except Exception as e:
            logger.error(f"Error downloading {url}: {e}")
    
    async def _process_downloaded_files(self):
        """Processa tutti i file scaricati."""
        for file_path in self.download_dir.iterdir():
            if file_path.is_file():
                try:
                    await self._process_downloaded_file(file_path)
                except Exception as e:
                    logger.error(f"Error processing {file_path}: {e}")
                    self.stats['errors'] += 1
    
    async def _process_downloaded_file(self, file_path: Path):
        """Processa un singolo file scaricato."""
        file_ext = file_path.suffix.lower()
        
        if file_ext == '.pdf':
            await self._process_pdf(file_path)
        elif file_ext in ['.xlsx', '.xls']:
            await self._process_excel(file_path)
        elif file_ext in ['.docx', '.doc']:
            await self._process_word(file_path)
    
    async def _process_pdf(self, file_path: Path):
        """Processa un file PDF."""
        if not PyPDF2:
            logger.warning("PyPDF2 not available for PDF processing")
            return
        
        try:
            content = ""
            
            # Try with pdfplumber first (better text extraction)
            if PDF:
                with PDF.open(file_path) as pdf:
                    for page in pdf.pages:
                        page_text = page.extract_text()
                        if page_text:
                            content += page_text + "\n"
            
            # Fallback to PyPDF2
            if not content:
                with open(file_path, 'rb') as file:
                    pdf_reader = PyPDF2.PdfReader(file)
                    for page in pdf_reader.pages:
                        content += page.extract_text() + "\n"
            
            if content:
                await self._create_document_from_file(file_path, content, "pdf")
                self.stats['pdfs_processed'] += 1
            
        except Exception as e:
            logger.error(f"Error processing PDF {file_path}: {e}")
    
    async def _process_excel(self, file_path: Path):
        """Processa un file Excel."""
        if not openpyxl:
            logger.warning("openpyxl not available for Excel processing")
            return
        
        try:
            workbook = load_workbook(file_path, read_only=True)
            content = ""
            
            for sheet_name in workbook.sheetnames:
                sheet = workbook[sheet_name]
                content += f"\\n--- Sheet: {sheet_name} ---\\n"
                
                for row in sheet.iter_rows(values_only=True):
                    row_text = " | ".join(str(cell) for cell in row if cell is not None)
                    if row_text.strip():
                        content += row_text + "\\n"
            
            if content:
                await self._create_document_from_file(file_path, content, "excel")
                self.stats['excel_processed'] += 1
            
        except Exception as e:
            logger.error(f"Error processing Excel {file_path}: {e}")
    
    async def _process_word(self, file_path: Path):
        """Processa un file Word."""
        if not docx:
            logger.warning("python-docx not available for Word processing")
            return
        
        try:
            doc = Document(file_path)
            content = ""
            
            for paragraph in doc.paragraphs:
                content += paragraph.text + "\\n"
            
            # Extract table content
            for table in doc.tables:
                for row in table.rows:
                    row_text = " | ".join(cell.text for cell in row.cells)
                    content += row_text + "\\n"
            
            if content:
                await self._create_document_from_file(file_path, content, "word")
                self.stats['word_processed'] += 1
            
        except Exception as e:
            logger.error(f"Error processing Word {file_path}: {e}")
    
    async def _create_document_from_file(self, file_path: Path, content: str, file_type: str):
        """Crea un documento estratto da un file."""
        content = self._clean_text(content)
        
        if len(content) < 50:  # Skip files with very little content
            return
        
        title = file_path.stem.replace('_', ' ').replace('-', ' ').title()
        
        # Extract keywords and categorize
        keywords, product_tags, category = self.keyword_extractor.extract_keywords(content, title)
        
        # Calculate confidence score
        confidence = self._calculate_confidence_score(content, title, keywords)
        
        # Create document
        doc = ExtractedDocument(
            id=self._generate_id(str(file_path)),
            content=content,
            title=title,
            source_url=f"file://{file_path}",
            file_type=file_type,
            category=category,
            keywords=keywords,
            product_tags=product_tags,
            technical_specs=self._extract_technical_specs(content),
            confidence_score=confidence,
            extraction_timestamp=datetime.now().isoformat(),
            file_hash=hashlib.md5(content.encode()).hexdigest(),
            size_bytes=len(content.encode()),
            language="en"
        )
        
        self.extracted_documents.append(doc)
        logger.info(f"Processed {file_type}: {title[:50]}... (Category: {category})")
    
    def _clean_text(self, text: str) -> str:
        """Pulisce e normalizza il testo."""
        # Remove extra whitespace
        text = re.sub(r'\\s+', ' ', text)
        # Remove special characters but keep technical specs
        text = re.sub(r'[^\\w\\s\\-\\.\\(\\)\\[\\]\\+\\*\\/\\%\\$\\&\\@\\#]', ' ', text)
        # Remove extra spaces
        text = ' '.join(text.split())
        return text.strip()
    
    def _extract_technical_specs(self, content: str) -> Dict[str, Any]:
        """Estrae specifiche tecniche strutturate."""
        specs = {}
        
        # Power specifications
        power_pattern = r'(\\d+)\\s*(?:w|watt|watts)'
        power_matches = re.findall(power_pattern, content, re.IGNORECASE)
        if power_matches:
            specs['power_watts'] = [int(p) for p in power_matches]
        
        # Frequency specifications
        freq_pattern = r'(\\d+)\\s*(?:hz|khz)'
        freq_matches = re.findall(freq_pattern, content, re.IGNORECASE)
        if freq_matches:
            specs['frequency_hz'] = freq_matches
        
        # SPL specifications
        spl_pattern = r'(\\d+)\\s*db'
        spl_matches = re.findall(spl_pattern, content, re.IGNORECASE)
        if spl_matches:
            specs['spl_db'] = [int(s) for s in spl_matches]
        
        return specs
    
    def _calculate_confidence_score(self, content: str, title: str, keywords: List[str]) -> float:
        """Calcola un punteggio di confidenza per la qualità del contenuto."""
        score = 0.5  # Base score
        
        # Content length bonus
        content_length = len(content)
        if content_length > 1000:
            score += 0.2
        elif content_length > 500:
            score += 0.1
        
        # Keywords bonus
        if len(keywords) > 5:
            score += 0.2
        elif len(keywords) > 2:
            score += 0.1
        
        # Title quality
        if title and len(title) > 10:
            score += 0.1
        
        # Technical content bonus
        technical_indicators = ['specification', 'datasheet', 'manual', 'guide', 'technical']
        if any(indicator in content.lower() for indicator in technical_indicators):
            score += 0.1
        
        return min(score, 1.0)
    
    def _generate_id(self, source: str) -> str:
        """Genera un ID univoco per il documento."""
        return hashlib.md5(f"{source}_{datetime.now().isoformat()}".encode()).hexdigest()
    
    def _validate_and_deduplicate(self):
        """Valida e rimuove duplicati dai documenti estratti."""
        # Remove duplicates based on content hash
        seen_hashes = set()
        unique_docs = []
        
        for doc in self.extracted_documents:
            if doc.file_hash not in seen_hashes:
                seen_hashes.add(doc.file_hash)
                unique_docs.append(doc)
            else:
                self.stats['duplicates_skipped'] += 1
        
        # Filter by confidence score
        high_quality_docs = [doc for doc in unique_docs if doc.confidence_score >= 0.6]
        
        self.extracted_documents = high_quality_docs
        logger.info(f"After validation: {len(self.extracted_documents)} high-quality documents")
    
    async def save_to_jsonl(self, output_path: Path):
        """Salva i documenti estratti in formato JSONL."""
        async with aiofiles.open(output_path, 'w', encoding='utf-8') as f:
            for doc in self.extracted_documents:
                await f.write(doc.to_jsonl() + '\\n')
        
        logger.info(f"Saved {len(self.extracted_documents)} documents to {output_path}")

# Main execution
async def main():
    """Funzione principale per eseguire l'estrazione."""
    
    # Verifica dependencies
    missing_deps = []
    if not PyPDF2:
        missing_deps.append("PyPDF2")
    if not openpyxl:
        missing_deps.append("openpyxl")
    if not docx:
        missing_deps.append("python-docx")
    if not BeautifulSoup:
        missing_deps.append("beautifulsoup4")
    
    if missing_deps:
        logger.warning(f"Missing dependencies: {missing_deps}")
        logger.info("Install with: pip install PyPDF2 openpyxl python-docx beautifulsoup4 pdfplumber")
    
    # Run extraction
    async with AdvancedExtractor() as extractor:
        documents = await extractor.extract_all()
        
        # Save results
        output_path = Path("karray_rag/data/karray_advanced_knowledge.jsonl")
        output_path.parent.mkdir(parents=True, exist_ok=True)
        await extractor.save_to_jsonl(output_path)
        
        # Print summary
        print(f"\\n=== EXTRACTION COMPLETE ===")
        print(f"Total documents: {len(documents)}")
        print(f"Categories found: {set(doc.category for doc in documents)}")
        print(f"Product tags found: {set(tag for doc in documents for tag in doc.product_tags)}")
        print(f"Average confidence: {sum(doc.confidence_score for doc in documents) / len(documents):.2f}")
        print(f"Statistics: {extractor.stats}")

if __name__ == "__main__":
    asyncio.run(main())